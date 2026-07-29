#!/usr/bin/env python3
"""
Mechanical Jurisprudence Document Extractor.

Extracts structured fields from Brazilian court documents (TJSP, TJMS, TJCE, TJRS)
using regex patterns, producing structured JSON ready for Qdrant ingestion.

Usage:
    cd /root/juris-search && source .venv/bin/activate
    python court_extractor.py --courts TJSP TJMS TJCE TJRS --max-per-court 5
    python court_extractor.py --courts TJSP --max-per-court 1  # single test
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import textwrap
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# ── Add .venv site-packages for PyPDF2 and python-docx ──────────────────────
_VENV_SP = "/home/disconzi1986_gmail_com/juris-search-VPS/.venv/lib/python3.12/site-packages"
if _VENV_SP not in sys.path:
    sys.path.insert(0, _VENV_SP)

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
try:
    from odf.opendocument import load as load_odt
    from odf.text import P
    HAS_ODF = True
except ImportError:
    HAS_ODF = False

# ── Reusable patterns from juris_indexer.py ──────────────────────────────────
CNJ_PROC_RE = re.compile(r"\b(\d{7}-?\d{2}\.?\d{4}\.?\d\.?\d{2}\.?\d{4})\b")
MONEY_RE = re.compile(r"R\$\s*[\d\.\,]+", re.IGNORECASE)
OUTCOME_PATTERNS: List[Tuple[str, re.Pattern]] = [
    ("negado_provimento",     re.compile(r"\bnegar(?:am)?\s+provimento\b", re.IGNORECASE)),
    ("dado_provimento",       re.compile(r"\bd(?:ar(?:am)?|eram)\s+provimento\b", re.IGNORECASE)),
    ("provimento_parcial",    re.compile(r"\bparcial(?:mente)?\s+provimento\b|\bprovimento\s+parcial\b", re.IGNORECASE)),
    ("reformada",             re.compile(r"\bsenten[cç]a\s+reformada\b", re.IGNORECASE)),
    ("mantida",               re.compile(r"\bmantida\s+a\s+senten[cç]a\b|\bmantenho\s+por\s+seus\s+pr[oó]prios\s+fundamentos\b", re.IGNORECASE)),
    ("procedente",            re.compile(r"\bjulgo\s+procedente\b|\bprocedente\s+o\s+pedido\b", re.IGNORECASE)),
    ("improcedente",          re.compile(r"\bjulgo\s+improcedente\b|\bimprocedente\s+o\s+pedido\b", re.IGNORECASE)),
    ("unanime",               re.compile(r"\bun[aâ]nime\b", re.IGNORECASE)),
]
EMENTA_END_RE = re.compile(r"\b(AC[ÓO]RD[ÃA]O|RELAT[ÓO]RIO|VOTOS?)\b")
TJRS_FNAME_RE = re.compile(
    r"inteiro_teor_(?P<numero>\d+)_(?P<ano>\d{4})_(?P<codigo>\d+)", re.IGNORECASE)
ESAJ_FNAME_RE = re.compile(
    r"(?:inteiro_teor|acordao|cdacordao)[_\-]?(?P<cdacordao>\d{6,})", re.IGNORECASE)

# ── Month name → number for Brazilian Portuguese ─────────────────────────────
_MONTHS = {
    "janeiro": 1, "fevereiro": 2, "março": 3, "marco": 3, "abril": 4,
    "maio": 5, "junho": 6, "julho": 7, "agosto": 8, "setembro": 9,
    "outubro": 10, "novembro": 11, "dezembro": 12,
}

# ── Legislation citation patterns ────────────────────────────────────────────
LEGISLACAO_RE = re.compile(
    r"(?:art(?:igo)?s?\s*[\d\.\-]+\s*(?:[e,]\s*[\d\.\-]+)*\s*(?:do|da|,)?\s*)?"
    r"(Lei\s+(?:[nN][º°]\.?\s*)?\d[\d\.]*/\d{4}|"
    r"Decreto[- ]Lei\s+(?:[nN][º°]\.?\s*)?\d[\d\.]*/\d{4}|"
    r"C[oó]digo\s+(?:de\s+)?(?:Processo\s+)?(?:Civil|Penal|Tribut[aá]rio|"
    r"de\s+Defesa\s+do\s+Consumidor|Eleitoral|Comercial|de\s+Tr[aâ]nsito\s+Brasileiro)|"
    r"Constitui[cç][aã]o\s+Federal)",
    re.IGNORECASE,
)

# ── Date patterns ────────────────────────────────────────────────────────────
DATE_LONG_RE = re.compile(
    r"(\d{1,2})\s+de\s+(janeiro|fevereiro|março|marco|abril|maio|junho|julho|"
    r"agosto|setembro|outubro|novembro|dezembro)\s+de\s+(\d{4})",
    re.IGNORECASE,
)
DATE_DDMMYYYY_RE = re.compile(r"\b(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})\b")


# ═══════════════════════════════════════════════════════════════════════════════
# Utilities
# ═══════════════════════════════════════════════════════════════════════════════

def _norm(text: str) -> str:
    """Normalize whitespace: collapse multiple spaces/newlines into single space."""
    return re.sub(r"\s+", " ", text).strip()


def _parse_date(text: str) -> Optional[str]:
    """Try multiple date formats, return ISO YYYY-MM-DD or None."""
    text = text.strip()
    # dd/mm/yyyy or dd-mm-yyyy
    m = DATE_DDMMYYYY_RE.search(text)
    if m:
        try:
            return f"{m.group(3)}-{int(m.group(2)):02d}-{int(m.group(1)):02d}"
        except ValueError:
            pass
    # "19 de maio de 2026"
    m = DATE_LONG_RE.search(text)
    if m:
        try:
            mon = _MONTHS.get(m.group(2).lower(), 0)
            if mon:
                return f"{m.group(3)}-{mon:02d}-{int(m.group(1)):02d}"
        except ValueError:
            pass
    return None


def _extract_outcomes(text: str) -> List[str]:
    seen = set()
    results = []
    for label, pat in OUTCOME_PATTERNS:
        if pat.search(text) and label not in seen:
            results.append(label)
            seen.add(label)
    return results


def _extract_cnj(text: str) -> Optional[str]:
    """Find first CNJ-format process number in text."""
    m = CNJ_PROC_RE.search(text)
    return m.group(1) if m else None


def _extract_money(text: str) -> List[str]:
    return MONEY_RE.findall(text)[:8]


def _extract_legislacao(text: str) -> List[str]:
    seen = set()
    result = []
    for m in LEGISLACAO_RE.finditer(text):
        item = _norm(m.group(0))
        if item.lower() not in seen:
            result.append(item)
            seen.add(item.lower())
    return result[:20]


def _classify(text: str, keywords: List[str]) -> List[str]:
    """Classify text into categories based on keyword presence."""
    found = []
    for kw in keywords:
        if re.search(r"\b" + re.escape(kw) + r"\b", text, re.IGNORECASE):
            found.append(kw)
    return found


def _load_master_index(path: str = "/home/disconzi1986_gmail_com/juris-search-VPS/master_index/master_index.json") -> dict:
    with open(path) as f:
        return json.load(f)


def _build_master_lookup(master: dict) -> Dict[str, dict]:
    """Build lookup by cdacordao and by raw_source_path basename."""
    by_cdacordao = {}
    by_rsp = {}
    for doc in master.get("documents", []):
        cd = doc.get("cdacordao")
        if cd:
            by_cdacordao[str(cd)] = doc
        rsp = doc.get("raw_source_path", "")
        if rsp:
            by_rsp[os.path.basename(rsp)] = doc
    return by_cdacordao, by_rsp


# ═══════════════════════════════════════════════════════════════════════════════
# Text Extraction
# ═══════════════════════════════════════════════════════════════════════════════

def extract_pdf_text(pdf_path: str) -> str:
    """Extract full text from PDF using PyPDF2."""
    reader = PdfReader(pdf_path)
    parts = []
    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            parts.append(txt)
    return "\n\n".join(parts)


def extract_docx_text(docx_path: str) -> str:
    """Extract full text from DOCX using python-docx."""
    doc = DocxDocument(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_doc_text(doc_path: str) -> str:
    """Extract full text from .doc file.
    
    First attempts to use LibreOffice to convert .doc to .docx,
    then extracts text from the resulting .docx.
    Falls back to basic text extraction if conversion fails.
    """
    try:
        # Try to use LibreOffice to convert .doc to .docx
        import subprocess
        import tempfile
        
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "converted.docx")
            try:
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "docx", 
                     "--outdir", tmpdir, doc_path],
                    timeout=30,
                    capture_output=True
                )
                if os.path.exists(output_path):
                    return extract_docx_text(output_path)
            except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
                pass  # Fall through to other methods
    except Exception:
        pass
    
    # Fallback: Try to extract using python-docx directly (works for some .doc files)
    try:
        doc = DocxDocument(doc_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        # If all else fails, return empty string
        return ""


def extract_html_text(html_path: str) -> str:
    """Extract readable text from HTML file using BeautifulSoup."""
    try:
        with open(html_path, "r", encoding="utf-8") as f:
            html = f.read()
    except UnicodeDecodeError:
        # Try with latin-1 if utf-8 fails
        with open(html_path, "r", encoding="latin-1") as f:
            html = f.read()
    
    soup = BeautifulSoup(html, "html.parser")
    
    # Remove script and style elements
    for tag in soup(["script", "style", "meta", "link"]):
        tag.decompose()
    
    # Get text and clean it up
    text = soup.get_text(separator="\n", strip=True)
    
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════════════
# Base Extractor
# ═══════════════════════════════════════════════════════════════════════════════

class BaseExtractor:
    tribunal: str = ""
    section_marker = re.compile(
        r"(I+\.?)\s*(?:-|–)?\s*(?:CASO\s+EM\s+EXAME|RELAT[ÓO]RIO)",
        re.IGNORECASE,
    )

    def __init__(self, text: str, filename: str, master_lookup: dict = None):
        self.raw_text = text
        self.text = _norm(text)
        self.filename = filename
        self.master = master_lookup or {}
        self.result: Dict[str, Any] = {
            "schema_version": 1,
            "extracted_at": datetime.now(timezone.utc).isoformat(),
            "source_file": filename,
            "tribunal": self.tribunal,
        }
        self.confidence: Dict[str, str] = {}

    def _set(self, key: str, value: Any, confidence: str = "high"):
        if value is not None and value != [] and value != "":
            self.result[key] = value
            self.confidence[key] = confidence

    def extract_all(self) -> dict:
        self._extract_common()
        self._extract_ementa()
        self._extract_outcomes()
        self._extract_legislacao()
        self._extract_assuntos()
        self._extract_court_specific()
        self.result["texto_inteiro"] = self.raw_text
        self.result["texto_length"] = len(self.raw_text)
        self.result["extraction_confidence"] = self.confidence
        return self.result

    def _extract_common(self):
        """Override in subclasses."""
        pass

    def _extract_ementa(self):
        """Extract ementa between EMENTA marker and next section (ACÓRDÃO/RELATÓRIO/VOTO)."""
        raw = self.raw_text
        # Find EMENTA start
        m_start = re.search(r"\bEMENTA\b", raw, re.IGNORECASE)
        if not m_start:
            m_start = re.search(r"\bE\s+M\s+E\s+N\s+T\s+A\b", raw, re.IGNORECASE)
        if m_start:
            rest = raw[m_start.end():]
            # Stop at next major section
            m_end = re.search(
                r"\b(?:A\s+C\s+[ÓO]\s+R\s+D\s+[ÃA]\s+O|AC[ÓO]RD[ÃA]O|RELAT[ÓO]RIO|VOTO\b|"
                r"I+\.?\s*(?:-|–)?\s*(?:CASO\s+EM\s+EXAME|RELAT[ÓO]RIO))",
                rest, re.IGNORECASE,
            )
            end_pos = m_end.start() if m_end else min(3000, len(rest))
            ementa = _norm(rest[:end_pos]).lstrip("-–: \t")
            if len(ementa) > 80:
                self._set("ementa", ementa, "high")
                return
        # Fallback to section-marker approach
        m = self.section_marker.search(self.text)
        if m:
            prefix = self.text[:m.start()]
            ementa_start = 0
            for marker in ["EMENTA", "E M E N T A", "Ementa"]:
                idx = prefix.rfind(marker)
                if idx > ementa_start:
                    ementa_start = idx + len(marker)
            if ementa_start > 0:
                ementa = _norm(prefix[ementa_start:]).lstrip("-–: \t")
                em_end = EMENTA_END_RE.search(ementa)
                if em_end:
                    ementa = ementa[:em_end.start()]
                if len(ementa) > 50:
                    self._set("ementa", ementa, "medium")
                    return
        # Last resort
        self._set("ementa", self.text[:2000], "low")

    def _extract_outcomes(self):
        outcomes = _extract_outcomes(self.text)
        self._set("outcome", outcomes, "high" if outcomes else "low")

    def _extract_legislacao(self):
        leg = _extract_legislacao(self.text)
        self._set("legislacao_citada", leg, "medium" if leg else "low")

    def _extract_assuntos(self):
        keywords = [
            "DIREITO PENAL", "DIREITO CIVIL", "DIREITO DO CONSUMIDOR",
            "DIREITO ADMINISTRATIVO", "DIREITO TRIBUTÁRIO", "DIREITO PROCESSUAL CIVIL",
            "DIREITO PROCESSUAL PENAL", "TRÁFICO DE DROGAS", "ASSOCIAÇÃO PARA O TRÁFICO",
            "ORGANIZAÇÃO CRIMINOSA", "HOMICÍDIO", "ROUBO", "FURTO", "ESTELIONATO",
            "LICITAÇÃO", "IMPROBIDADE ADMINISTRATIVA", "RESPONSABILIDADE CIVIL",
            "DANO MORAL", "ALIMENTOS", "FAMÍLIA", "EXECUÇÃO FISCAL",
            "CRIMES DE RESPONSABILIDADE", "CONTRAVENÇÃO PENAL", "VIOLÊNCIA DOMÉSTICA",
            "POSSE DE ARMA", "RECEPTAÇÃO", "LAVAGEM DE DINHEIRO",
        ]
        found = _classify(self.text, keywords)
        self._set("assuntos", found, "high" if found else "low")

    def _extract_court_specific(self):
        """Override in subclasses."""
        pass


# ═══════════════════════════════════════════════════════════════════════════════
# TJSP Extractor
# ═══════════════════════════════════════════════════════════════════════════════

class TJSPExtractor(BaseExtractor):
    tribunal = "TJSP"

    def _extract_common(self):
        text = self.text

        # Registro
        m = re.search(r"Registro:\s*(\d{4}\.\d+)", text)
        if m:
            self.result.setdefault("court_specific", {})["registro"] = m.group(1)

        # Numero processo
        m = re.search(
            r"(?:Apelação|Agravo|Habeas\s+Corpus)\s+\w+\s+(?:[nN][º°]\.?\s*)?(\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4})",
            text,
        )
        if m:
            self._set("numero_processo", m.group(1))
        else:
            cnj = _extract_cnj(text)
            if cnj:
                self._set("numero_processo", cnj, "medium")

        # Classe
        m = re.search(r"(Apelação\s+Criminal|Apelação\s+C[íi]vel|Habeas\s+Corpus|Agravo\s+de\s+Instrumento|Agravo\s+em\s+Execu[cç][aã]o)", text, re.IGNORECASE)
        if m:
            self._set("classe", m.group(1))

        # Órgão julgador
        m = re.search(r"(\d+ª\s+C[aâ]mara\s+de\s+Direito\s+\w+)", text)
        if m:
            self._set("orgao_julgador", m.group(1))
            self.result.setdefault("court_specific", {})["camara"] = m.group(1)

        # Comarca
        m = re.search(r"Comarca\s+de\s+(.+?)(?:,|\.|\n|\s+em\s+que)", text)
        if m:
            self._set("comarca", _norm(m.group(1)))

        # Relator — after ACÓRDÃO block, name before "Relator"
        m = re.search(r"([A-ZÀ-Ú][A-ZÀ-Ú\s]{3,40})\s*\n\s*Relator", self.raw_text)
        if m:
            self._set("relator", _norm(m.group(1)))

        # Data julgamento
        m = re.search(r"São Paulo,\s*(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})", text)
        if m:
            dt = _parse_date(m.group(1))
            self._set("data_julgamento", dt)

        # Partes
        partes = {}
        for role, label in [("apelantes", "Apelante"), ("apelados", "Apelad[ao]")]:
            pat = re.compile(rf"{label}s?:?\s*(.+?)(?=\n\n|\n(?:Apelad|Ju[ií]zo|Advogad|EMENTA|DIREITO))", re.IGNORECASE | re.DOTALL)
            m = pat.search(self.raw_text)
            if m:
                names = re.split(r"\s*[,;]\s*|\s+e\s+", _norm(m.group(1)))
                names = [n for n in names if len(n) > 3]
                if names:
                    partes[role] = names
        if partes:
            self._set("partes", partes)

        # Decisão
        m = re.search(r'"[^"]*((?:deram|negaram|proveram|acolheram|rejeitaram)[^"]*)"', text, re.IGNORECASE)
        if m:
            self._set("decisao", _norm(m.group(0).strip('"')))
        else:
            m = re.search(r"(Deram\s+provimento(?:\s+parcial)?.*?V\.\s*U\.)", text, re.IGNORECASE)
            if m:
                self._set("decisao", _norm(m.group(1)), "medium")

        # Voto número
        m = re.search(r"Voto\s+(?:[nN][º°]\.?\s*)?(\d+(?:\.\d+)?)", text)
        if m:
            self.result.setdefault("court_specific", {})["voto_numero"] = m.group(1)

        # Votação
        m = re.search(r"\b(V\.\s*U\.|POR\s+MAIORIA|POR\s+UNANIMIDADE)\b", text, re.IGNORECASE)
        if m:
            self._set("votacao", _norm(m.group(1)))

    def _extract_ementa(self):
        """TJSP ementa is between second header block (APTE/APDO) and 'Trata-se de'."""
        raw = self.raw_text
        # Find second "PODER JUDICIÁRIO" after "Assinatura Eletrônica"
        ass_idx = raw.find("Assinatura")
        if ass_idx > 0:
            rest = raw[ass_idx:]
            # Skip past APTE/APDO block, find ementa start
            m_apte = re.search(r"APTE\.\s*:", rest)
            m_apdo = re.search(r"APDO\.\s*:", rest)
            if m_apte or m_apdo:
                # Start after APDO block (the later of the two)
                ementa_start = max(m_apte.end() if m_apte else 0, m_apdo.end() if m_apdo else 0)
                ementa_text = rest[ementa_start:]
                # Find end at "Trata-se de" or "Vistos" or "I." section
                m_end = re.search(r"\b(Trata-se\s+de|Vistos\s*[,;]|I+\.?\s*(?:-|–)?\s*(?:CASO|RELAT))", ementa_text, re.IGNORECASE)
                end_pos = m_end.start() if m_end else min(3000, len(ementa_text))
                ementa = _norm(ementa_text[:end_pos]).lstrip("*–: \t")
                if len(ementa) > 80:
                    self._set("ementa", ementa, "high")
                    return
        # Fallback to base class
        super()._extract_ementa()

    def _extract_court_specific(self):
        pass  # Done inside _extract_common


# ═══════════════════════════════════════════════════════════════════════════════
# TJMS Extractor
# ═══════════════════════════════════════════════════════════════════════════════

class TJMSExtractor(BaseExtractor):
    tribunal = "TJMS"

    def _extract_common(self):
        text = self.text
        raw = self.raw_text

        # Número processo
        m = re.search(
            r"(?:Apelação|Agravo)\s+(?:C[íi]vel|Criminal)\s+-\s+(?:[nN][º°]\.?\s*)?(\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4})",
            text,
        )
        if m:
            self._set("numero_processo", m.group(1))
        else:
            cnj = _extract_cnj(text)
            if cnj:
                self._set("numero_processo", cnj, "medium")

        # Comarca — use raw_text to stop at newline
        m = re.search(
            r"(?:[nN][º°]\.?\s*)?\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}\s*-\s*(.+?)(?:\n|$)",
            raw[:2000],
        )
        if m:
            self._set("comarca", _norm(m.group(1)))

        # Câmara — first few lines
        m = re.search(r"(\d+ª\s+C[aâ]mara\s+\w+)", raw[:500])
        if m:
            self._set("orgao_julgador", m.group(1))
            self.result.setdefault("court_specific", {})["camara"] = m.group(1)

        # Relator — use raw_text to stop at newline
        m = re.search(
            r"^Relator\s+(?:Designad[oa]\s*)?(?:–|-)?\s*(?:Ex[º°]\.?\s*Sr\.?\s*Des\.?\s*)?(.+?)$",
            raw[:2000],
            re.IGNORECASE | re.MULTILINE,
        )
        if m:
            self._set("relator", _norm(m.group(1)))

        # Data sessão
        m = re.search(
            r"Tribunal\s+de\s+Justi[cç]a\s+do\s+Estado\s+de\s+Mato\s+Grosso\s+do\s+Sul\s*\n\s*(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})",
            self.raw_text,
        )
        if m:
            dt = _parse_date(m.group(1))
            self._set("data_sessao", dt)

        # Partes
        partes = {}
        for role, label in [("apelantes", "Apelante"), ("apelados", "Apelad[oa]")]:
            pat = re.compile(
                rf"^{label}\s*:\s*(.+?)$",
                re.IGNORECASE | re.MULTILINE,
            )
            names = []
            for m in pat.finditer(self.raw_text):
                name = _norm(m.group(1))
                if name and len(name) > 3:
                    names.append(re.sub(r"\s*\(OAB:.*$", "", name))
            if names:
                partes[role] = names
        if partes:
            self._set("partes", partes)

        # Advogados + OAB
        advs = []
        for m in re.finditer(
            r"Advogad[oa]\s*:\s*(.+?)\s*\(OAB:\s*(\d+/\w+)\)",
            self.raw_text,
        ):
            advs.append({"nome": _norm(m.group(1)), "oab": m.group(2)})
        if advs:
            self._set("advogados", advs)
            self.result.setdefault("court_specific", {})["oab_advogados"] = [
                f"{a['nome']} (OAB: {a['oab']})" for a in advs
            ]

        # Promotor
        m = re.search(r"Prom\.\s*Justi[cç]a\s*:\s*(.+?)(?:\n|$)", self.raw_text, re.IGNORECASE)
        if m:
            self.result.setdefault("court_specific", {})["promotor"] = _norm(m.group(1))

        # Interessado
        interessados = []
        for m in re.finditer(r"Interessad[oa]\s*:\s*(.+?)(?:\n|$)", self.raw_text, re.IGNORECASE):
            interessados.append(_norm(m.group(1)))
        if interessados:
            self.result.setdefault("court_specific", {})["interessados"] = interessados

        # Decisão from D E C I S Ã O section
        dec_text = ""
        idx = text.find("D E C I S Ã O")
        if idx < 0:
            idx = text.find("DECISÃO")
        if idx >= 0:
            dec_text = text[idx:idx + 1500]
        else:
            dec_text = text[-2000:]
        m = re.search(
            r"(?:POR\s+(MAIORIA|UNANIMIDADE),\s*)?(DERAM|NEGARAM)\s+PROVIMENTO(?:\s+PARCIAL)?",
            dec_text,
            re.IGNORECASE,
        )
        if m:
            self._set("decisao", _norm(m.group(0)))

        # Votação detalhe
        vot_detail = []
        if re.search(r"VENCIDO\s+O\s+RELATOR", dec_text, re.IGNORECASE):
            vot_detail.append("vencido_relator")
        m = re.search(r"NOS\s+TERMOS\s+DO\s+VOTO\s+DO\s+(\d+[º°]\s+VOGA?L)", dec_text, re.IGNORECASE)
        if m:
            vot_detail.append(m.group(1))
        m = re.search(r"\b(POR\s+MAIORIA|POR\s+UNANIMIDADE)\b", dec_text, re.IGNORECASE)
        if m:
            self._set("votacao", _norm(m.group(1)))
            vot_detail.append(_norm(m.group(1)))
        if vot_detail:
            self.result.setdefault("court_specific", {})["votacao_detalhe"] = vot_detail

        # Data julgamento — from footer
        m = re.search(r"(?:Campo\s+Grande|Data)[,\s]*(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})", text[-2000:])
        if m:
            dt = _parse_date(m.group(1))
            self._set("data_julgamento", dt)

    def _extract_ementa(self):
        """TJMS ementa is between EMENTA marker and ACÓRDÃO marker."""
        raw = self.raw_text
        m_start = re.search(r"\bEMENTA\b", raw, re.IGNORECASE)
        if m_start:
            rest = raw[m_start.end():]
            m_end = re.search(r"\bA\s+C\s+[ÓO]\s+R\s+D\s+[ÃA]\s+O\b|AC[ÓO]RD[ÃA]O", rest, re.IGNORECASE)
            end_pos = m_end.start() if m_end else min(3000, len(rest))
            ementa = _norm(rest[:end_pos]).lstrip("-– \t")
            if len(ementa) > 80:
                self._set("ementa", ementa, "high")
                return
        # Fallback
        self._set("ementa", self.text[:2000], "low")

    def _extract_court_specific(self):
        pass  # Done inside _extract_common


# ═══════════════════════════════════════════════════════════════════════════════
# TJCE Extractor
# ═══════════════════════════════════════════════════════════════════════════════

class TJCEExtractor(BaseExtractor):
    tribunal = "TJCE"

    def _extract_common(self):
        text = self.text

        # Gabinete
        m = re.search(r"GABINETE\s+DESEMBARGADOR\s+(.+)", text[:500])
        if m:
            self.result.setdefault("court_specific", {})["gabinete"] = _norm(m.group(1))

        # Número processo
        m = re.search(r"Processo:\s*(\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4})", text[:1000])
        if m:
            self._set("numero_processo", m.group(1))
        else:
            cnj = _extract_cnj(text)
            if cnj:
                self._set("numero_processo", cnj, "medium")

        # Classe
        m = re.search(r"Processo:\s*\d+-\d+\s*-\s*(.+?)(?:\n|$)", text[:1000])
        if m:
            self._set("classe", _norm(m.group(1)))

        # Partes
        partes = {}
        for role, label in [("apelantes", "Apelantes?"), ("apelados", "Apelados?")]:
            pat = re.compile(
                rf"{label}\s*:\s*(.+?)(?:\.\s*(?:Apelad|Corr[eé]u|\.\s*$)|$)",
                re.IGNORECASE | re.DOTALL,
            )
            m = pat.search(self.raw_text[:3000])
            if m:
                names = re.split(r"\s*[,;]\s*|\s+e\s+", _norm(m.group(1)))
                names = [n for n in names if len(n) > 3]
                if names:
                    partes[role] = names
        if partes:
            self._set("partes", partes)

        # Corréu
        m = re.search(r"Corr[eé]u\s*:\s*(.+?)(?:\n|$)", self.raw_text[:3000], re.IGNORECASE)
        if m:
            self.result.setdefault("court_specific", {})["correu"] = _norm(m.group(1))

        # Relator — at the end of the document
        m = re.search(r"([A-ZÀ-Ú][A-ZÀ-Ú\s]{3,40})\s*\n\s*(?:Desembargador|Relator)", self.raw_text[-2000:])
        if m:
            self._set("relator", _norm(m.group(1)))
        if not self.result.get("relator"):
            m = re.search(r"(?:DESEMBARGADOR|Desembargador)\s+(.+)", text[-3000:])
            if m:
                self._set("relator", _norm(m.group(1)), "medium")

        # Decisão — near end
        dec_text = text[-5000:]
        m = re.search(
            r"(?:DERAM|NEGARAM|DERAM\s+PARCIAL\s+PROVIMENTO|NEGARAM\s+PROVIMENTO).*?(?:\.|\n)",
            dec_text,
            re.IGNORECASE,
        )
        if m:
            self._set("decisao", _norm(m.group(0)))

        # Votação
        m = re.search(r"\b(POR\s+MAIORIA|POR\s+UNANIMIDADE|V\.\s*U\.)\b", dec_text, re.IGNORECASE)
        if m:
            self._set("votacao", _norm(m.group(1)))

        # Data julgamento
        m = re.search(r"Fortaleza,\s*(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})", text[-3000:])
        if m:
            dt = _parse_date(m.group(1))
            self._set("data_julgamento", dt)

    def _extract_court_specific(self):
        pass


# ═══════════════════════════════════════════════════════════════════════════════
# TJRS Extractor
# ═══════════════════════════════════════════════════════════════════════════════

class TJRSExtractor(BaseExtractor):
    tribunal = "TJRS"

    def _extract_common(self):
        text = self.text
        raw = self.raw_text

        # Número processo from filename
        m = TJRS_FNAME_RE.search(self.filename)
        if m:
            self._set("numero_processo", m.group("numero"))
            self._set("ano", m.group("ano"))
            self._set("codigo", m.group("codigo"))

        # Classe from first line
        m = re.search(r"^(?:APELA[ÇC][ÃA]O|APELAÇÕES)\s+(\w+)", text, re.IGNORECASE)
        if m:
            self._set("classe", f"Apelação {m.group(1).capitalize()}")

        # Cross-reference with master index for metadata not in body
        if self.master:
            self._cross_reference_master()

        # Preliminares section
        p_start = raw.upper().find("PRELIMINARES")
        if p_start >= 0:
            p_text = raw[p_start:]
            m_end = p_text.upper().find("MÉRITO")
            if m_end > 200:
                p_text = p_text[:m_end]
            preliminares = [s.strip() for s in re.split(r"\n\s*\n", p_text) if len(s.strip()) > 50]
            self.result.setdefault("court_specific", {})["preliminares"] = preliminares[:30]

        # Fatos
        fatos = []
        for m in re.finditer(r"(\d+)[º°]\s+FATO\.\s*([^\.]+)\.?\s*", raw, re.IGNORECASE):
            fato_num = int(m.group(1))
            crime = _norm(m.group(2))
            # Get more context after the match
            ctx_start = m.end()
            next_fato = re.search(r"\d+[º°]\s+FATO\.", raw[ctx_start:ctx_start + 5000], re.IGNORECASE)
            ctx_end = ctx_start + next_fato.start() if next_fato else ctx_start + 2000
            contexto = _norm(raw[ctx_start:ctx_end])[:500]
            fatos.append({"numero": fato_num, "crime": crime, "contexto": contexto})
        if fatos:
            self.result.setdefault("court_specific", {})["fatos"] = fatos
            self.result.setdefault("court_specific", {})["quantidade_fatos"] = len(fatos)

        # Dosimetria
        dos_text = ""
        d_start = raw.upper().find("DOSIMETRIA")
        if d_start >= 0:
            dos_text = raw[d_start:]
        if dos_text:
            reus = []
            for m in re.finditer(
                r"([A-ZÀ-Ú][A-ZÀ-Ú\s]{5,40}(?:DOS\s+SANTOS|DE\s+\w+|DA\s+\w+|J[UÚ]NIOR)?)\."
                r"(.*?)"
                r"(?:(?:fixo|estabele[cç]o|reduz[oi]\s*(?:a|para))\s*(?:a\s*)?(?:pena|san[cç][aã]o)\s*(?:em|para|definitiva)?\s*"
                r"(\d+)\s*(?:a\s*)?(?:ano|ANOS?)(?:[^.]*?(?:(\d+)\s*(?:m[êe]s|MESES?))?[^.]*?(?:(\d+)\s*(?:dia|DIAS?))?)?"
                r"[^.]*?regime\s+inicial\s+(fechado|semiaberto|aberto))",
                dos_text,
                re.IGNORECASE | re.DOTALL,
            ):
                nome = _norm(m.group(1))
                pena_anos = int(m.group(3)) if m.group(3) else 0
                pena_meses = int(m.group(4)) if m.group(4) else 0
                pena_dias = int(m.group(5)) if m.group(5) else 0
                regime = m.group(6).lower() if m.group(6) else ""
                reus.append({
                    "nome": nome,
                    "pena_anos": pena_anos,
                    "pena_meses": pena_meses,
                    "pena_dias": pena_dias,
                    "regime": regime,
                })
            if reus:
                self.result.setdefault("court_specific", {})["dosimetria"] = reus
                self.result.setdefault("court_specific", {})["quantidade_reus"] = len(reus)

        # Decisão — from final paragraph
        m = re.search(
            r"(?:DERAM|NEGARAM|DERAM\s+PARCIAL)\s+PROVIMENTO.*?UN[AÂ]NIME",
            text[-5000:],
            re.IGNORECASE,
        )
        if m:
            self._set("decisao", _norm(m.group(0)), "medium")

        # Votação
        m = re.search(r"\bUN[AÂ]NIME\b", text[-2000:], re.IGNORECASE)
        if m:
            self._set("votacao", "UNANIME")

    def _cross_reference_master(self):
        """Fill missing metadata from master index."""
        if not self.master:
            return
        # Find this document in master lookup by cdacordao or filename
        master_doc = None
        for key in self.master:
            doc = self.master[key]
            rsp = doc.get("raw_source_path", "")
            if self.filename in rsp or os.path.basename(rsp) in self.filename:
                master_doc = doc
                break
        if not master_doc:
            return
        for field in ["relator", "comarca", "data_julgamento", "data_publicacao",
                       "orgao_julgador", "assunto", "classe"]:
            val = master_doc.get(field)
            if val and not self.result.get(field):
                self._set(field, val, "cross_reference")

    def _extract_court_specific(self):
        pass  # Done inside _extract_common


# ═══════════════════════════════════════════════════════════════════════════════
# Extractor registry
# ═══════════════════════════════════════════════════════════════════════════════

EXTRACTORS: Dict[str, type] = {
    "TJSP": TJSPExtractor,
    "TJMS": TJMSExtractor,
    "TJCE": TJCEExtractor,
    "TJRS": TJRSExtractor,
}


# ═══════════════════════════════════════════════════════════════════════════════
# Main Pipeline
# ═══════════════════════════════════════════════════════════════════════════════

def _find_files_for_courts(
    base_dir: str,
    courts: List[str],
    master_lookup_cd: dict,
    master_lookup_rsp: dict,
) -> Dict[str, List[str]]:
    """Find files in base_dir/PDF and base_dir/docx, mapping to courts."""
    base = Path(base_dir)
    pdf_dir = base / "PDF"
    docx_dir = base / "docx"

    court_files: Dict[str, List[str]] = {c: [] for c in courts}

    # Scan PDFs
    if pdf_dir.exists() and any(c in courts for c in ["TJSP", "TJMS", "TJCE"]):
        for f in sorted(pdf_dir.glob("*.pdf")):
            cd_match = ESAJ_FNAME_RE.search(f.name)
            if cd_match:
                cdacordao = cd_match.group("cdacordao")
                master_doc = master_lookup_cd.get(cdacordao)
                if master_doc:
                    trib = master_doc.get("tribunal", "")
                    if trib in courts:
                        court_files[trib].append(str(f))
                        continue
            # Fallback: try raw_source_path lookup
            master_doc = master_lookup_rsp.get(f.name)
            if master_doc:
                trib = master_doc.get("tribunal", "")
                if trib in courts:
                    court_files[trib].append(str(f))

    # Scan DOCX (always TJRS)
    if docx_dir.exists() and "TJRS" in courts:
        for f in sorted(docx_dir.glob("*.docx")):
            court_files["TJRS"].append(str(f))

    return court_files


def process_file(filepath: str, tribunal: str, master_lookup_rsp: dict) -> Optional[dict]:
    """Process a single file: extract text, run court extractor, return result dict."""
    ext = os.path.splitext(filepath)[1].lower()
    fname = os.path.basename(filepath)

    # Extract text
    if ext == ".pdf":
        text = extract_pdf_text(filepath)
    elif ext == ".docx":
        text = extract_docx_text(filepath)
    elif ext == ".doc":
        text = extract_doc_text(filepath)
    elif ext == ".html" or ext == ".htm":
        text = extract_html_text(filepath)
    else:
        print(f"  SKIP {fname}: unsupported format {ext}")
        return None

    if not text or len(text) < 100:
        print(f"  SKIP {fname}: text too short ({len(text)} chars)")
        return None

    # Run extractor
    extractor_cls = EXTRACTORS.get(tribunal)
    if not extractor_cls:
        print(f"  SKIP {fname}: no extractor for {tribunal}")
        return None

    extractor = extractor_cls(text, fname, master_lookup_rsp)
    result = extractor.extract_all()

    return result


def main():
    parser = argparse.ArgumentParser(description="Mechanical Jurisprudence Document Extractor")
    parser.add_argument("--courts", nargs="+", default=["TJSP", "TJMS", "TJCE", "TJRS"],
                        choices=["TJSP", "TJMS", "TJCE", "TJRS"],
                        help="Courts to process")
    parser.add_argument("--max-per-court", type=int, default=None,
                        help="Max documents per court (for testing)")
    parser.add_argument("--input-dir", default="/home/disconzi1986_gmail_com/juris-search-VPS/court_samples/jurisprudence-documents",
                        help="Directory with PDF/ and docx/ subdirectories")
    parser.add_argument("--output-dir", 
                        default=os.environ.get(
                            "JURIS_SEARCH_EXTRACTIONS_DIR",
                            os.path.join(os.path.dirname(os.path.abspath(__file__)), "extracted_documents")
                        ),
                        help="Output directory for extracted JSONs")
    parser.add_argument("--master-index", default="/home/disconzi1986_gmail_com/juris-search-VPS/master_index/master_index.json",
                        help="Path to master_index.json for cross-referencing")
    args = parser.parse_args()

    # Load master index
    master = _load_master_index(args.master_index)
    cd_lookup, rsp_lookup = _build_master_lookup(master)

    # Find files
    print(f"Scanning {args.input_dir} ...")
    court_files = _find_files_for_courts(args.input_dir, args.courts, cd_lookup, rsp_lookup)
    for trib, files in court_files.items():
        print(f"  {trib}: {len(files)} candidates")

    # Limit per court
    if args.max_per_court:
        for trib in court_files:
            court_files[trib] = court_files[trib][:args.max_per_court]

    # Process
    os.makedirs(args.output_dir, exist_ok=True)
    total = 0
    field_counts: Counter = Counter()

    for tribunal in args.courts:
        files = court_files.get(tribunal, [])
        if not files:
            print(f"\n{tribunal}: no files found")
            continue
        print(f"\n{'='*60}")
        print(f"{tribunal}: processing {len(files)} files")
        print(f"{'='*60}")

        for i, filepath in enumerate(files, 1):
            fname = os.path.basename(filepath)
            print(f"  [{i}/{len(files)}] {fname} ...", end=" ", flush=True)
            result = process_file(filepath, tribunal, rsp_lookup)
            if result is None:
                print("FAILED")
                continue

            # Count extracted fields
            extracted = [k for k, v in result.items()
                         if k not in ("schema_version", "extracted_at", "source_file",
                                      "tribunal", "texto_inteiro", "texto_length",
                                      "extraction_confidence", "court_specific")
                         and v is not None and v != [] and v != ""]
            for f in extracted:
                field_counts[f] += 1

            # Write output
            doc_id = os.path.splitext(fname)[0]
            out_path = os.path.join(args.output_dir, f"{tribunal}_{doc_id}.json")
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(result, f, ensure_ascii=False, indent=2)

            conf_items = [f"{k}={v}" for k, v in result.get("extraction_confidence", {}).items()
                          if v == "low"]
            conf_str = f"  LOW_CONF: {', '.join(conf_items[:5])}" if conf_items else ""
            print(f"OK ({len(extracted)} fields){conf_str}")
            total += 1

    # Summary
    print(f"\n{'='*60}")
    print(f"SUMMARY: {total} documents processed")
    print(f"Output: {args.output_dir}")
    print(f"\nField extraction rates:")
    for field, count in field_counts.most_common():
        pct = count / max(total, 1) * 100
        print(f"  {field}: {count}/{total} ({pct:.0f}%)")


# ═══════════════════════════════════════════════════════════════════════════════
# Automation API — callable from post-download chain
# ═══════════════════════════════════════════════════════════════════════════════

import uuid as _uuid
import urllib.request
import urllib.error

_QDRANT_API_BASE = os.environ.get("JURIS_SEARCH_QDRANT_API", "http://localhost:8066")
_QDRANT_COLLECTION = os.environ.get("JURIS_SEARCH_QDRANT_COLLECTION", "juris_br_v1")
_EXTRACTIONS_DIR = os.environ.get(
    "JURIS_SEARCH_EXTRACTIONS_DIR",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "extracted_documents"),
)


def _doc_uuid(doc_id: str) -> str:
    return str(_uuid.uuid5(_uuid.NAMESPACE_DNS, f"juris-search:{doc_id}"))


def _http_post_json(url: str, payload: dict, timeout: int = 120) -> Tuple[int, Any]:
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            body = resp.read().decode("utf-8")
            return resp.status, json.loads(body) if body else {}
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8") if e.fp else ""
        try:
            return e.code, json.loads(body) if body else {"detail": str(e)}
        except json.JSONDecodeError:
            return e.code, {"detail": body[:500]}


def _build_embedding_text(doc: dict) -> str:
    """Compose text blob for semantic embedding."""
    parts = [f"Tribunal: {doc.get('tribunal') or 'N/D'}"]
    proc = doc.get("numero_processo") or doc.get("cnj_numero") or "N/D"
    parts.append(f"Processo: {proc}")
    if doc.get("classe"):
        parts.append(f"Classe: {doc['classe']}")
    if doc.get("relator"):
        parts.append(f"Relator: {doc['relator']}")
    if doc.get("orgao_julgador"):
        parts.append(f"Orgao julgador: {doc['orgao_julgador']}")
    if doc.get("comarca"):
        parts.append(f"Comarca: {doc['comarca']}")
    if doc.get("data_julgamento"):
        parts.append(f"Julgado em: {doc['data_julgamento']}")
    outcome = doc.get("outcome")
    if outcome:
        parts.append(f"Resultado: {', '.join(outcome) if isinstance(outcome, list) else outcome}")
    assuntos = doc.get("assuntos")
    if assuntos:
        parts.append(f"Assuntos: {', '.join(assuntos) if isinstance(assuntos, list) else assuntos}")
    header = " | ".join(parts)
    body_parts = []
    if doc.get("ementa"):
        body_parts.append(doc["ementa"])
    if doc.get("decisao"):
        body_parts.append(doc["decisao"])
    legislacao = doc.get("legislacao_citada")
    if legislacao:
        body_parts.append("Legislacao: " + (", ".join(legislacao) if isinstance(legislacao, list) else legislacao))
    body = "\n\n".join(body_parts) if body_parts else ""
    return f"{header}\n\n{body}" if body else header


def extract_file(file_path: str, tribunal: str) -> Optional[dict]:
    """Extract structured fields from a single downloaded document.

    Args:
        file_path: Path to the .pdf or .docx file.
        tribunal: One of TJSP, TJMS, TJCE, TJRS.

    Returns:
        Extracted document dict, or None on failure.
    """
    ext = os.path.splitext(file_path)[1].lower()
    fname = os.path.basename(file_path)

    if ext == ".pdf":
        text = extract_pdf_text(file_path)
    elif ext == ".docx":
        text = extract_docx_text(file_path)
    elif ext == ".doc":
        text = extract_doc_text(file_path)
    elif ext == ".html" or ext == ".htm":
        text = extract_html_text(file_path)
    else:
        print(f"[extract] SKIP {fname}: unsupported format {ext}")
        return None

    if not text or len(text) < 100:
        print(f"[extract] SKIP {fname}: text too short ({len(text)} chars)")
        return None

    extractor_cls = EXTRACTORS.get(tribunal)
    if not extractor_cls:
        print(f"[extract] SKIP {fname}: no extractor for tribunal {tribunal}")
        return None

    # Build minimal master lookup from existing master_index.json
    rsp_lookup = {}
    master_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "master_index", "master_index.json")
    try:
        if os.path.exists(master_path):
            master = _load_master_index(master_path)
            _, rsp_lookup = _build_master_lookup(master)
    except Exception:
        pass

    extractor = extractor_cls(text, fname, rsp_lookup)
    result = extractor.extract_all()

    if not result:
        print(f"[extract] SKIP {fname}: extraction returned empty")
        return None

    # Save to extracted_documents/
    os.makedirs(_EXTRACTIONS_DIR, exist_ok=True)
    proc = result.get("numero_processo") or os.path.splitext(fname)[0]
    safe_id = re.sub(r"[^\w\-\.]", "_", proc)[:120]
    out_path = os.path.join(_EXTRACTIONS_DIR, f"{tribunal}_{safe_id}.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"[extract] OK {fname} → {os.path.basename(out_path)}")

    return result


def ingest_extracted_to_qdrant(
    doc: dict,
    qdrant_api: str = _QDRANT_API_BASE,
    collection: str = _QDRANT_COLLECTION,
) -> Dict[str, Any]:
    """Ingest a single extracted document into Qdrant.

    Delegates to ingest_to_qdrant.ingest_single for correct payload format.
    """
    try:
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        import ingest_to_qdrant as _ingest
        return _ingest.ingest_single(doc, collection=collection, api_base=qdrant_api)
    except Exception as exc:
        return {"ok": False, "error": f"import failed: {exc}"}



# ═══════════════════════════════════════════════════════════════════════════════
# Programmatic API — importable from other modules
# ═══════════════════════════════════════════════════════════════════════════════

# Use environment variable with fallback to project-relative path
_EXTR_PATH = os.environ.get(
    "JURIS_SEARCH_EXTRACTIONS_DIR",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "extracted_documents")
)
EXTRACTIONS_DIR = Path(_EXTR_PATH)
MASTER_INDEX_PATH = Path("/home/disconzi1986_gmail_com/juris-search-VPS/master_index/master_index.json")
QDRANT_API_DEFAULT = "http://localhost:8114"
QDRANT_COLLECTION = "juris_br_v1"


def extract_and_ingest(filepath: str, tribunal: str) -> Dict[str, Any]:
    """
    Combined: extract structured fields + ingest to Qdrant.
    Called by routes_download._auto_extract_and_ingest().
    Returns {"ok": bool, "proc": str, "error": str}.
    """
    extracted = extract_file(filepath, tribunal)
    if extracted is None:
        return {"ok": False, "proc": os.path.basename(filepath), "error": "extraction failed"}
    proc = extracted.get("numero_processo") or os.path.basename(filepath)
    result = ingest_extracted_to_qdrant(extracted)
    if result.get("ok"):
        return {"ok": True, "proc": proc, "point_id": result.get("point_id")}
    return {"ok": False, "proc": proc, "error": result.get("error", "unknown")}


def _load_master_lookup() -> dict:
    """Load master_index.json and build rsp_lookup for extractor context."""
    if not MASTER_INDEX_PATH.is_file():
        return {}
    try:
        master = json.loads(MASTER_INDEX_PATH.read_text(encoding="utf-8"))
        _, rsp = _build_master_lookup(master)
        return rsp
    except Exception:
        return {}


def extract_file(filepath: str, tribunal: str) -> Optional[dict]:
    """
    Extract structured fields from a single downloaded document.
    Writes JSON to extracted_documents/ and returns the result dict.
    Returns None on failure.
    """
    master_lookup = _load_master_lookup()
    result = process_file(filepath, tribunal, master_lookup)
    if result is None:
        return None

    # Write output file
    EXTRACTIONS_DIR.mkdir(parents=True, exist_ok=True)
    fname = os.path.basename(filepath)
    doc_id = os.path.splitext(fname)[0]
    out_path = EXTRACTIONS_DIR / f"{tribunal}_{doc_id}.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    return result


if __name__ == "__main__":
    main()
