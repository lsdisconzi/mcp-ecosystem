"""File text extraction utilities for uploads and storage pipeline."""

import os
import sys
import re
import io as _io
import base64
import shutil
import tempfile
import subprocess
import logging
import mimetypes
from pathlib import Path
from typing import Optional, Tuple, Dict, Any, List

from modules.config import (
    DOWNLOADS_BASE_DIR,
    DOCX_JURISPRUDENCE_DIR,
    DOCX_NORMALIZE_FOR_COMPAT,
)

logger = logging.getLogger("juris-search.file_extraction")

# ── Optional dependencies ───────────────────────────────────────────────────

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


# ── LibreOffice discovery ───────────────────────────────────────────────────

def _find_libreoffice_binary() -> Optional[str]:
    candidates = []
    if sys.platform == "darwin":
        candidates.extend([
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            shutil.which("libreoffice"),
            shutil.which("soffice"),
        ])
    elif sys.platform == "win32":
        candidates.extend([
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            shutil.which("soffice"),
            shutil.which("libreoffice"),
        ])
    else:
        candidates.extend([
            shutil.which("libreoffice"),
            shutil.which("soffice"),
        ])

    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate
    return None


# ── Path-based extractors (used by storage pipeline) ────────────────────────

def _extract_text_from_docx_path(path: Path) -> str:
    if not DocxDocument:
        raise RuntimeError("python-docx is not installed")

    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs)


def _extract_text_from_pdf_path(path: Path) -> str:
    if not PdfReader and not pdfplumber:
        raise RuntimeError("PyPDF2 e pdfplumber não estão instalados")

    chunks: List[str] = []
    if PdfReader:
        with open(path, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages[:200]:
                text = page.extract_text() or ""
                text = text.strip()
                if text:
                    chunks.append(text)

    # Fallback / supplement with pdfplumber when PyPDF2 yields little text
    # (e.g. PDFs with unsupported CMaps such as Big5 fonts).
    if pdfplumber and (not chunks or sum(len(c) for c in chunks) < 200):
        try:
            with pdfplumber.open(str(path)) as pdf:
                plumber_chunks: List[str] = []
                for page in pdf.pages[:200]:
                    text = page.extract_text() or ""
                    text = text.strip()
                    if text:
                        plumber_chunks.append(text)
                if plumber_chunks:
                    chunks = plumber_chunks
        except Exception as exc:
            logger.debug("pdfplumber fallback failed for %s: %s", path, exc)

    return "\n\n".join(chunks)


def _extract_text_from_html_bytes(raw: bytes) -> Tuple[str, str]:
    html_text = raw.decode("utf-8", errors="replace")

    if BeautifulSoup:
        soup = BeautifulSoup(html_text, "html.parser")
        extracted = soup.get_text("\n", strip=True)
        cleaned = re.sub(r"\n{3,}", "\n\n", extracted).strip()
        return cleaned, "html_bs4"

    extracted = re.sub(r"<script[\s\S]*?</script>", " ", html_text, flags=re.IGNORECASE)
    extracted = re.sub(r"<style[\s\S]*?</style>", " ", extracted, flags=re.IGNORECASE)
    extracted = re.sub(r"<[^>]+>", " ", extracted)
    extracted = re.sub(r"\s+", " ", extracted).strip()
    return extracted, "html_regex"


# ── Upload-facing extractors ────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not PdfReader and not pdfplumber:
        return "[PyPDF2 e pdfplumber não instalados — não foi possível extrair texto do PDF]"

    pages: List[str] = []
    if PdfReader:
        try:
            reader = PdfReader(_io.BytesIO(file_bytes))
            for page in reader.pages[:20]:
                t = page.extract_text()
                if t:
                    pages.append(t)
        except Exception as e:
            logger.debug("PyPDF2 extraction failed: %s", e)

    # Fallback to pdfplumber when PyPDF2 produced no usable text
    if pdfplumber and not pages:
        try:
            with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages[:20]:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
        except Exception as e:
            logger.debug("pdfplumber extraction failed: %s", e)

    if not pages:
        return "[Não foi possível extrair texto do PDF]"
    return "\n\n".join(pages)[:8000]


def extract_text_from_docx(file_bytes: bytes) -> str:
    if not DocxDocument:
        return "[python-docx não instalado]"
    try:
        doc = DocxDocument(_io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)[:8000]
    except Exception as e:
        return f"[Erro ao ler DOCX: {e}]"


def extract_text_from_image(file_bytes: bytes) -> str:
    """For images, return base64 to send to DeepSeek vision."""
    return base64.b64encode(file_bytes).decode("utf-8")


def process_uploaded_file(filename: str, content: bytes) -> Dict[str, Any]:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    result = {"filename": filename, "type": ext, "text": None, "image_b64": None}

    if ext == "pdf":
        result["text"] = extract_text_from_pdf(content)
    elif ext in ("docx", "doc"):
        result["text"] = extract_text_from_docx(content)
    elif ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp"):
        result["image_b64"] = extract_text_from_image(content)
    else:
        try:
            result["text"] = content.decode("utf-8", errors="replace")[:8000]
        except Exception:
            result["text"] = "[Formato não suportado]"

    return result
